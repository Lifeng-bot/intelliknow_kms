
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import markdown

def create_docx(md_file_path, output_path):
    """Create a DOCX file from a Markdown file."""
    # Read markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Create a new Document
    doc = Document()

    # Parse markdown content
    lines = md_content.split('\n')
    for line in lines:
        if line.startswith('# '):
            # Heading 1
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            # Heading 2
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # Heading 3
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            # Bullet point
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip() == '':
            # Empty line
            doc.add_paragraph()
        else:
            # Regular paragraph
            doc.add_paragraph(line)

    # Save the document
    doc.save(output_path)
    print(f"Created DOCX: {output_path}")

def create_pdf(md_file_path, output_path):
    """Create a PDF file from a Markdown file."""
    # Read markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Create PDF
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()

    # Add custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=RGBColor(31, 119, 180),
        spaceAfter=30,
        alignment=1  # Center alignment
    )

    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=RGBColor(31, 119, 180),
        spaceAfter=12
    )

    heading3_style = ParagraphStyle(
        'CustomHeading3',
        parent=styles['Heading3'],
        fontSize=12,
        textColor=RGBColor(31, 119, 180),
        spaceAfter=10
    )

    # Parse markdown content
    lines = md_content.split('\n')
    for line in lines:
        if line.startswith('# '):
            # Heading 1
            story.append(Paragraph(line[2:], title_style))
            story.append(Spacer(1, 0.2*inch))
        elif line.startswith('## '):
            # Heading 2
            story.append(Paragraph(line[3:], heading2_style))
            story.append(Spacer(1, 0.1*inch))
        elif line.startswith('### '):
            # Heading 3
            story.append(Paragraph(line[4:], heading3_style))
            story.append(Spacer(1, 0.1*inch))
        elif line.startswith('- '):
            # Bullet point
            story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
        elif line.strip() != '':
            # Regular paragraph
            story.append(Paragraph(line, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))

    # Build PDF
    doc.build(story)
    print(f"Created PDF: {output_path}")

def main():
    # Define input and output directories
    input_dir = 'data/sample_documents'
    output_dir = 'data/sample_documents'

    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)

    # Process each markdown file
    md_files = [
        'Financial_Policy.md',
        'HR_Policy.md',
        'Legal_Policy.md'
    ]

    for md_file in md_files:
        md_path = os.path.join(input_dir, md_file)
        base_name = os.path.splitext(md_file)[0]

        # Create DOCX
        docx_path = os.path.join(output_dir, f"{base_name}.docx")
        create_docx(md_path, docx_path)

        # Create PDF
        pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
        create_pdf(md_path, pdf_path)

if __name__ == "__main__":
    main()
